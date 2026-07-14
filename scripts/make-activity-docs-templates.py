#!/usr/bin/env python3
"""
make-activity-docs-templates.py
Convert KCnotice raw templates into docxtemplater-compatible templates.
Run from keichi project root: python3 scripts/make-activity-docs-templates.py
"""
import os
from docx import Document

SRC = os.path.join(os.path.dirname(__file__), '..', 'public', 'templates')
OUT = SRC

def consolidate_runs(para, new_text, preserve_run_idx=0):
    runs = para.runs
    if not runs:
        para.add_run(new_text)
        return
    base = runs[min(preserve_run_idx, len(runs) - 1)]
    base.text = new_text
    for r in runs:
        if r is not base:
            r.text = ''

def set_cell_single_run(cell, new_text):
    for para in cell.paragraphs:
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ''
            return
    cell.paragraphs[0].add_run(new_text)

def patch_notice_header(paras):
    p0 = paras[0]
    runs = p0.runs
    if len(runs) >= 4:
        runs[3].text = '家長通告{noticeNum}'
        for r in runs[4:]:
            r.text = ''

def make_notice_t1(src_path, out_path):
    doc = Document(src_path)
    paras = doc.paragraphs
    patch_notice_header(paras)
    consolidate_runs(paras[1], '【{activityName}】')
    consolidate_runs(paras[5], '{bodyText}')
    consolidate_runs(paras[8], '{contactLine}')
    consolidate_runs(paras[17], '{issueDateCn}')
    table = doc.tables[0]
    set_cell_single_run(table.rows[0].cells[1], '{activityName}')
    set_cell_single_run(table.rows[1].cells[1], '{sessionDate}')
    set_cell_single_run(table.rows[2].cells[1], '{sessionTime}')
    set_cell_single_run(table.rows[3].cells[1], '{sessionLocation}')
    set_cell_single_run(table.rows[4].cells[1], '{teacherName}老師')
    doc.save(out_path)
    print(f'  saved: {os.path.basename(out_path)}')

def make_notice_t4(src_path, out_path):
    doc = Document(src_path)
    paras = doc.paragraphs
    patch_notice_header(paras)
    consolidate_runs(paras[1], '【{activityName}】')
    consolidate_runs(paras[5], '{bodyText}')
    consolidate_runs(paras[8], '{contactLine}')
    consolidate_runs(paras[17], '{issueDateCn}')
    table = doc.tables[0]
    set_cell_single_run(table.rows[0].cells[1], '{activityName}')
    set_cell_single_run(table.rows[1].cells[1], '{sessionDates}')
    set_cell_single_run(table.rows[2].cells[1], '{sessionTime}')
    set_cell_single_run(table.rows[3].cells[1], '{sessionLocation}')
    set_cell_single_run(table.rows[4].cells[1], '{teacherName}老師')
    doc.save(out_path)
    print(f'  saved: {os.path.basename(out_path)}')

def make_notice_t2(src_path, out_path):
    doc = Document(src_path)
    paras = doc.paragraphs
    patch_notice_header(paras)
    consolidate_runs(paras[1], '【{activityName}】')
    consolidate_runs(paras[5], '{bodyText}')
    consolidate_runs(paras[8], '{contactLine}')
    consolidate_runs(paras[17], '{issueDateCn}')
    table = doc.tables[0]
    set_cell_single_run(table.rows[0].cells[1], '{sessionAName}')
    if len(table.rows[0].cells) > 2:
        set_cell_single_run(table.rows[0].cells[2], '{sessionBName}')
    set_cell_single_run(table.rows[1].cells[1], '{sessionADate}')
    if len(table.rows[1].cells) > 2:
        set_cell_single_run(table.rows[1].cells[2], '{sessionBDate}')
    set_cell_single_run(table.rows[2].cells[1], '{sessionATime}')
    if len(table.rows[2].cells) > 2:
        set_cell_single_run(table.rows[2].cells[2], '{sessionBTime}')
    set_cell_single_run(table.rows[3].cells[1], '{sessionALocation}')
    if len(table.rows[3].cells) > 2:
        set_cell_single_run(table.rows[3].cells[2], '{sessionBLocation}')
    set_cell_single_run(table.rows[4].cells[1], '{teacherName}老師')
    if len(table.rows[4].cells) > 2:
        set_cell_single_run(table.rows[4].cells[2], '{teacherName}老師')
    doc.save(out_path)
    print(f'  saved: {os.path.basename(out_path)}')

def make_tutor_signin(src_path, out_path):
    doc = Document(src_path)
    paras = doc.paragraphs
    consolidate_runs(paras[3], '導師姓名/ 機構名稱：{tutorName}')
    consolidate_runs(paras[4], '          活動名稱：{activityName}')
    table = doc.tables[0]
    print(f'    導師簽到 table rows: {len(table.rows)}')
    MAX_SESSIONS = min(len(table.rows) - 1, 15)
    for i in range(MAX_SESSIONS):
        row = table.rows[i + 1]
        n = i + 1
        cells = row.cells
        if len(cells) > 0:
            set_cell_single_run(cells[0], '{s%dDate}' % n)
        if len(cells) > 1:
            set_cell_single_run(cells[1], '{s%dArrive}' % n)
        if len(cells) > 3:
            set_cell_single_run(cells[3], '{s%dLeave}' % n)
    doc.save(out_path)
    print(f'  saved: {os.path.basename(out_path)}')

if __name__ == '__main__':
    print('Converting KCnotice templates → docxtemplater placeholders...\n')
    make_notice_t1(os.path.join(SRC,'通告範本1.docx'), os.path.join(OUT,'activity-notice-t1.docx'))
    make_notice_t4(os.path.join(SRC,'通告範本4.docx'), os.path.join(OUT,'activity-notice-t4.docx'))
    make_notice_t2(os.path.join(SRC,'通告範本2.docx'), os.path.join(OUT,'activity-notice-t2.docx'))
    make_tutor_signin(os.path.join(SRC,'導師簽到.docx'), os.path.join(OUT,'activity-tutor-signin.docx'))
    print('\nDone.')
