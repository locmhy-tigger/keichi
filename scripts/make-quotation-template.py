#!/usr/bin/env python3
"""
make_template.py — Convert KCquotation/template_base.docx into a docxtemplater
template, written to keichi/public/templates/quotation.docx.

We insert {placeholder} tags into the EXACT run positions verified by dumping
the document (see dump.py / dump2.py / dump3.py / dump4.py). Run formatting
(esp. Wingdings 2 for checkboxes) is preserved because we only change .text.

Mapping is authoritative; assertions guard against structural drift.
"""
import sys
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SRC = "/Users/kenc/Documents/Dev/CCC_Kei_Chi/KCquotation/template_base.docx"
OUT = "/Users/kenc/Documents/Dev/CCC_Kei_Chi/keichi/public/templates/quotation.docx"

doc = Document(SRC)
placed = []  # log every tag we place

def set_run(para, idx, text):
    para.runs[idx].text = text

def clear_after(para, idx):
    for r in para.runs[idx + 1:]:
        r.text = ""

def put_tag_para(para, tag):
    """Replace a paragraph's entire text with a single {tag}, keeping run[0] format."""
    if para.runs:
        para.runs[0].text = tag
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(tag)

def tag_first_run(para, tag):
    """Put {tag} in run[0], clear the rest (for multi-run value cells)."""
    assert para.runs, f"no runs in para for {tag}"
    para.runs[0].text = tag
    clear_after(para, 0)

def set_vmerge_continue(cell):
    tcPr = cell._tc.tcPr
    if tcPr is None:
        from docx.oxml import OxmlElement
        tcPr = OxmlElement('w:tcPr')
        cell._tc.insert(0, tcPr)
    vm = tcPr.find(qn('w:vMerge'))
    if vm is None:
        from docx.oxml import OxmlElement
        vm = OxmlElement('w:vMerge')
        tcPr.append(vm)
    # remove val attribute => "continue"
    for a in list(vm.attrib):
        del vm.attrib[a]

def clear_cell_text(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ""

# ── para[3] header line (assert 45 runs) ────────────────────────────────────
p3 = doc.paragraphs[3]
assert len(p3.runs) == 45, f"para[3] has {len(p3.runs)} runs, expected 45"
set_run(p3, 1, "{quotationDate}"); placed.append("quotationDate")
# method checkboxes (Wingdings 2 runs)
for idx, name in [(5, "methodPhoneBox"), (8, "methodFaxBox"), (11, "methodMailBox"), (14, "methodOtherBox")]:
    set_run(p3, idx, "{" + name + "}"); placed.append(name)
# recommended supplier name (runs 27,28,29)
set_run(p3, 27, "{recommendedSupplier}"); set_run(p3, 28, ""); set_run(p3, 29, "")
placed.append("recommendedSupplier")
# price lower/higher checkboxes (Wingdings runs 35, 39)
set_run(p3, 35, "{priceLowerBox}"); placed.append("priceLowerBox")
set_run(p3, 39, "{priceHigherBox}"); placed.append("priceHigherBox")
# higher price reason (run 43)
set_run(p3, 43, "{higherPriceReason}"); placed.append("higherPriceReason")

# ── para[5] body line (assert 66 runs) ──────────────────────────────────────
p5 = doc.paragraphs[5]
assert len(p5.runs) == 66, f"para[5] has {len(p5.runs)} runs, expected 66"
# fewer suppliers reason (runs 5,6)
set_run(p5, 5, "{fewerSuppliersReason}"); set_run(p5, 6, ""); placed.append("fewerSuppliersReason")
# category checkboxes (Wingdings 16,19,22)
for idx, name in [(16, "catFixedBox"), (19, "catConsumableBox"), (22, "catOtherBox")]:
    set_run(p5, idx, "{" + name + "}"); placed.append(name)
# category other (runs 25,26,27)
set_run(p5, 25, "{categoryOther}"); set_run(p5, 26, ""); set_run(p5, 27, ""); placed.append("categoryOther")
# department (runs 30,31,32)
set_run(p5, 30, "{department}"); set_run(p5, 31, ""); set_run(p5, 32, ""); placed.append("department")
# purpose (runs 35,36,37)
set_run(p5, 35, "{purpose}"); set_run(p5, 36, ""); set_run(p5, 37, ""); placed.append("purpose")
# delivery date (run 50)
set_run(p5, 50, "{deliveryDate}"); placed.append("deliveryDate")
# funding source (runs 57,58)
set_run(p5, 57, "{fundingSource}"); set_run(p5, 58, ""); placed.append("fundingSource")

# ── signature fill paragraphs (locate by label, take next para) ─────────────
def find_fill_para(label):
    for i, p in enumerate(doc.paragraphs[:-1]):
        if label in p.text:
            return doc.paragraphs[i + 1]
    raise AssertionError(f"label {label!r} not found")

req = find_fill_para("索取報價人")
assert len(req.runs) == 7, f"requestor fill para has {len(req.runs)} runs"
set_run(req, 0, "{requestorName}"); placed.append("requestorName")
set_run(req, 2, "{requestorRank}"); placed.append("requestorRank")
set_run(req, 6, "{requestorDate}"); placed.append("requestorDate")

dh = find_fill_para("科組負責人")
assert len(dh.runs) == 7, f"depthead fill para has {len(dh.runs)} runs"
set_run(dh, 0, "{deptHeadName}"); placed.append("deptHeadName")
set_run(dh, 2, "{deptHeadRank}"); placed.append("deptHeadRank")
set_run(dh, 6, "{deptHeadDate}"); placed.append("deptHeadDate")

# ── table[0] ────────────────────────────────────────────────────────────────
# IMPORTANT: operate on raw <w:tr>/<w:tc> XML elements by index. python-docx's
# row.cells resolution ALIASES the vertically-merged shared columns (tc 0,1,2)
# across rows to the same object, so mutating via .cells corrupts row 1.
# Raw element access touches each distinct tc unambiguously.
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

tbl = doc.tables[0]._tbl
trs = tbl.findall(qn("w:tr"))
assert len(trs) == 4, f"table has {len(trs)} rows, expected 4"

def tc(row_idx, col_idx):
    return trs[row_idx].findall(qn("w:tc"))[col_idx]

def paras_of(tc_el):
    return [Paragraph(p, tc_el) for p in tc_el.findall(qn("w:p"))]

def put_tag(tc_el, para_idx, tag):
    """Replace paragraph's full text with {tag} (keeps run[0] format), or add run."""
    p = paras_of(tc_el)[para_idx]
    if p.runs:
        p.runs[0].text = tag
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(tag)

def tag_run0(tc_el, para_idx, tag):
    p = paras_of(tc_el)[para_idx]
    assert p.runs, f"no runs in para[{para_idx}] for {tag}"
    p.runs[0].text = tag
    for r in p.runs[1:]:
        r.text = ""

def clear_tc(tc_el):
    for p in paras_of(tc_el):
        for r in p.runs:
            r.text = ""

def vmerge_continue(tc_el):
    tcPr = tc_el.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc_el.insert(0, tcPr)
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        vm = OxmlElement("w:vMerge")
        tcPr.append(vm)
    for a in list(vm.attrib):
        del vm.attrib[a]

# Row 1 — items (shared cols, vMerge=restart kept) + supplier A
put_tag(tc(1, 1), 0, "{quotationName}"); placed.append("quotationName")
put_tag(tc(1, 1), 1, "{item1Name}"); placed.append("item1Name")
put_tag(tc(1, 1), 2, "{item2Name}"); placed.append("item2Name")
put_tag(tc(1, 1), 3, "{item3Name}"); placed.append("item3Name")
put_tag(tc(1, 2), 1, "{item1Qty}"); placed.append("item1Qty")
put_tag(tc(1, 2), 2, "{item2Qty}"); placed.append("item2Qty")
put_tag(tc(1, 2), 3, "{item3Qty}"); placed.append("item3Qty")
tag_run0(tc(1, 3), 0, "{supAName}"); placed.append("supAName")
tag_run0(tc(1, 3), 1, "{supATel}"); placed.append("supATel")
tag_run0(tc(1, 4), 0, "{supAPrices}"); placed.append("supAPrices")
tag_run0(tc(1, 5), 0, "{supATotal}"); placed.append("supATotal")
pa = paras_of(tc(1, 6))[0]
if pa.runs:
    pa.runs[0].text = "{supAAdopt}"
    for r in pa.runs[1:]:
        r.text = ""
else:
    pa.add_run("{supAAdopt}")
placed.append("supAAdopt")

# Row 2 — merge shared cols (0,1,2) into row 1 + supplier B
for col in (0, 1, 2):
    vmerge_continue(tc(2, col))
    clear_tc(tc(2, col))
tag_run0(tc(2, 3), 0, "{supBName}"); placed.append("supBName")
tag_run0(tc(2, 3), 1, "{supBTel}"); placed.append("supBTel")
tag_run0(tc(2, 4), 0, "{supBPrices}"); placed.append("supBPrices")
tag_run0(tc(2, 5), 0, "{supBTotal}"); placed.append("supBTotal")
pb = paras_of(tc(2, 6))[0]
if pb.runs:
    pb.runs[0].text = "{supBAdopt}"
    for r in pb.runs[1:]:
        r.text = ""
else:
    pb.add_run("{supBAdopt}")
placed.append("supBAdopt")

# Row 3 — optional 3rd supplier, unused: merge shared cols + clear supplier cells
for col in (0, 1, 2):
    vmerge_continue(tc(3, col))
    clear_tc(tc(3, col))
for col in (3, 4, 5, 6):
    clear_tc(tc(3, col))

doc.save(OUT)

print(f"Saved -> {OUT}")
print(f"Tags placed ({len(placed)}):")
for t in placed:
    print(f"  - {t}")
